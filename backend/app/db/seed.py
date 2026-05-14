from __future__ import annotations

import hashlib
from datetime import UTC, datetime, timedelta
from decimal import Decimal
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from sqlalchemy import select

from app.core.config import settings
from app.core.security import get_password_hash
from app.db.session import SessionLocal
from app.models.audit_log import AuditLog
from app.models.document import Document
from app.models.review import Review
from app.models.risk_finding import RiskFinding
from app.models.user import User

DEMO_USERS = [
    ("client@example.com", "client"),
    ("reviewer@example.com", "reviewer"),
    ("admin@example.com", "admin"),
]


def seed_users() -> None:
    with SessionLocal() as db:
        for email, role in DEMO_USERS:
            existing = db.scalar(select(User).where(User.email == email))
            if existing is not None:
                continue
            db.add(
                User(
                    email=email,
                    role=role,
                    password_hash=get_password_hash("password123"),
                )
            )
        db.commit()


def seed_sample_documents() -> None:
    with SessionLocal() as db:
        users = {user.email: user for user in db.scalars(select(User)).all()}
        client = users.get("client@example.com")
        reviewer = users.get("reviewer@example.com")
        admin = users.get("admin@example.com")
        if client is None or reviewer is None or admin is None:
            return

        now = datetime.now(UTC)
        for index, spec in enumerate(_sample_document_specs(now), start=1):
            existing = db.scalar(
                select(Document).where(
                    Document.user_id == client.id,
                    Document.filename == spec["filename"],
                )
            )
            if existing is not None:
                _ensure_docx_file(Path(existing.storage_path), spec["paragraphs"])
                continue

            uploaded_at = now - timedelta(hours=24 - index)
            processed_at = uploaded_at + timedelta(minutes=3) if spec["processing_status"] != "processing" else None
            storage_path, size_bytes, sha256 = _write_sample_file(
                user_id=str(client.id),
                filename=spec["filename"],
                paragraphs=spec["paragraphs"],
            )
            document = Document(
                user_id=client.id,
                filename=spec["filename"],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                size_bytes=size_bytes,
                sha256=sha256,
                storage_path=str(storage_path),
                processing_status=spec["processing_status"],
                review_status=spec["review_status"],
                classification=spec["classification"],
                classification_confidence=Decimal(spec["classification_confidence"]),
                summary=spec["summary"],
                extracted_text="\n\n".join(spec["paragraphs"]),
                risk_score=spec["risk_score"],
                ai_confidence=Decimal(spec["ai_confidence"]),
                flag_reasons=spec["flag_reasons"],
                expiry_date=spec.get("expiry_date"),
                uploaded_at=uploaded_at,
                processed_at=processed_at,
            )
            db.add(document)
            db.flush()

            for finding in spec["findings"]:
                db.add(
                    RiskFinding(
                        document_id=document.id,
                        rule_code=finding["rule_code"],
                        severity=finding["severity"],
                        snippet=finding["snippet"],
                        suggestion=finding["suggestion"],
                    )
                )

            _add_seed_audit_logs(
                db,
                document=document,
                actor_id=client.id,
                processed_at=processed_at,
                review_status=spec["review_status"],
            )

            if spec["review_status"] in {"admin_approved", "admin_rejected"}:
                db.add(
                    Review(
                        document_id=document.id,
                        reviewer_id=admin.id if spec["review_status"] == "admin_approved" else reviewer.id,
                        decision=spec["review_status"],
                        comment=spec["review_comment"],
                        override_ai=True,
                    )
                )
                db.add(
                    AuditLog(
                        action="admin.decision.submitted",
                        target_type="document",
                        target_id=document.id,
                        actor_id=admin.id if spec["review_status"] == "admin_approved" else reviewer.id,
                        payload={
                            "decision": spec["review_status"],
                            "comment": spec["review_comment"],
                            "seeded": True,
                        },
                    )
                )

        db.commit()


def _sample_document_specs(now: datetime) -> list[dict[str, Any]]:
    return [
        {
            "filename": "demo-safe-contract.docx",
            "processing_status": "completed",
            "review_status": "ai_approved",
            "classification": "contract",
            "classification_confidence": "0.9200",
            "summary": "Tài liệu hợp đồng có vẻ rủi ro thấp, chất lượng trích xuất được đánh giá là tốt.",
            "risk_score": 8,
            "ai_confidence": "0.8800",
            "flag_reasons": [],
            "expiry_date": (now + timedelta(days=45)).date(),
            "findings": [],
            "paragraphs": [
                "Master Services Agreement between Client Co and Vendor Co with clear governing law, termination rights, signature blocks, and payment terms.",
                "The agreement includes a 30-day termination notice, venue selection, and authorized signatory language.",
            ],
        },
        {
            "filename": "demo-safe-nda.docx",
            "processing_status": "completed",
            "review_status": "ai_approved",
            "classification": "nda",
            "classification_confidence": "0.8400",
            "summary": "Tài liệu NDA có vẻ rủi ro thấp, chất lượng trích xuất được đánh giá là tốt.",
            "risk_score": 12,
            "ai_confidence": "0.8100",
            "flag_reasons": [],
            "expiry_date": (now + timedelta(days=120)).date(),
            "findings": [],
            "paragraphs": [
                "Mutual Non-Disclosure Agreement covering confidential information, recipient obligations, disclosing party rights, governing law, and signatures.",
                "The confidentiality obligations are limited to three years and include standard exclusions.",
            ],
        },
        {
            "filename": "demo-high-value-contract.docx",
            "processing_status": "completed",
            "review_status": "pending_admin",
            "classification": "contract",
            "classification_confidence": "0.7100",
            "summary": "Tài liệu hợp đồng bị gắn cờ với điểm rủi ro 100. Vấn đề chính: Thiếu chữ ký, Giá trị cao, Thiếu luật điều chỉnh.",
            "risk_score": 100,
            "ai_confidence": "0.6800",
            "flag_reasons": ["MISSING_SIGNATURE", "HIGH_VALUE", "NO_GOVERNING_LAW"],
            "expiry_date": (now + timedelta(days=18)).date(),
            "findings": [
                _finding("MISSING_SIGNATURE", "high"),
                _finding("HIGH_VALUE", "critical", "Detected amount 250,000.00."),
                _finding("NO_GOVERNING_LAW", "medium"),
            ],
            "paragraphs": [
                "Enterprise contract between Client Co and Strategic Vendor for USD 250,000 in services.",
                "The agreement includes commercial obligations but omits governing law and signature language.",
            ],
        },
        {
            "filename": "demo-expiring-renewal.docx",
            "processing_status": "completed",
            "review_status": "pending_admin",
            "classification": "contract",
            "classification_confidence": "0.6600",
            "summary": "Tài liệu hợp đồng bị gắn cờ với điểm rủi ro 55. Vấn đề chính: Sắp hết hạn, Tự động gia hạn.",
            "risk_score": 55,
            "ai_confidence": "0.6400",
            "flag_reasons": ["EXPIRY_SOON", "AUTO_RENEWAL"],
            "expiry_date": (now + timedelta(days=5)).date(),
            "findings": [
                _finding("EXPIRY_SOON", "medium", f"Document appears to expire on {(now + timedelta(days=5)).date().isoformat()}"),
                _finding("AUTO_RENEWAL", "medium"),
            ],
            "paragraphs": [
                "Vendor agreement with automatic renewal language and an expiration date within the next week.",
                "The contract automatically renews unless either party gives notice before the renewal window closes.",
            ],
        },
        {
            "filename": "demo-approved-after-review.docx",
            "processing_status": "completed",
            "review_status": "admin_approved",
            "classification": "nda",
            "classification_confidence": "0.7600",
            "summary": "Tài liệu NDA đã được người rà soát duyệt sau khi kiểm tra các điều khoản bị gắn cờ.",
            "risk_score": 48,
            "ai_confidence": "0.7000",
            "flag_reasons": ["NO_TERMINATION_CLAUSE"],
            "expiry_date": (now + timedelta(days=35)).date(),
            "findings": [_finding("NO_TERMINATION_CLAUSE", "high")],
            "review_comment": "Đã kiểm tra phụ lục, điều khoản chấm dứt nằm ở phụ lục B.",
            "paragraphs": [
                "NDA for a pilot project with confidentiality obligations and a governing law clause.",
                "The automated rules missed a termination reference in the appendix.",
            ],
        },
        {
            "filename": "demo-rejected-invoice.docx",
            "processing_status": "completed",
            "review_status": "admin_rejected",
            "classification": "invoice",
            "classification_confidence": "0.7300",
            "summary": "Tài liệu hóa đơn bị từ chối do giá trị cao và thiếu thông tin xác nhận.",
            "risk_score": 82,
            "ai_confidence": "0.6900",
            "flag_reasons": ["HIGH_VALUE", "CONFIDENCE_LOW"],
            "expiry_date": None,
            "findings": [
                _finding("HIGH_VALUE", "critical", "Detected amount 180,000.00."),
                _finding("CONFIDENCE_LOW", "medium", "Classification confidence is 0.52."),
            ],
            "review_comment": "Từ chối vì hóa đơn thiếu PO và cần xác minh thủ công với finance.",
            "paragraphs": [
                "Invoice total USD 180,000 with payment due date and tax details.",
                "The supporting purchase order is not attached.",
            ],
        },
        {
            "filename": "demo-processing-upload.docx",
            "processing_status": "processing",
            "review_status": "processing",
            "classification": None,
            "classification_confidence": "0.0000",
            "summary": "Tài liệu đang chờ pipeline xử lý.",
            "risk_score": 0,
            "ai_confidence": "0.0000",
            "flag_reasons": [],
            "expiry_date": None,
            "findings": [],
            "paragraphs": [
                "Draft policy uploaded for processing.",
                "This seeded record demonstrates the in-progress state in the client document list.",
            ],
        },
        {
            "filename": "demo-failed-extraction.docx",
            "processing_status": "failed",
            "review_status": "failed",
            "classification": "unknown",
            "classification_confidence": "0.2500",
            "summary": "Không thể xử lý tài liệu mẫu này.",
            "risk_score": 0,
            "ai_confidence": "0.0000",
            "flag_reasons": ["LOW_EXTRACTION_QUALITY"],
            "expiry_date": None,
            "findings": [_finding("LOW_EXTRACTION_QUALITY", "high")],
            "paragraphs": [
                "Unreadable scan placeholder.",
                "This seeded record demonstrates the failed-processing state.",
            ],
        },
    ]


def _finding(rule_code: str, severity: str, snippet: str | None = None) -> dict[str, str | None]:
    suggestions = {
        "MISSING_SIGNATURE": "Confirm the agreement includes signature and execution details.",
        "HIGH_VALUE": "Escalate high-value documents for manual approval.",
        "EXPIRY_SOON": "Review renewal or extension terms before the expiry window closes.",
        "NO_TERMINATION_CLAUSE": "Add explicit termination rights and notice periods.",
        "NO_GOVERNING_LAW": "Specify the governing law and venue for dispute resolution.",
        "AUTO_RENEWAL": "Confirm renewal notice periods and opt-out conditions.",
        "CONFIDENCE_LOW": "Review the document manually before trusting the automated verdict.",
        "LOW_EXTRACTION_QUALITY": "Check the source document quality or upload a text-based version.",
    }
    snippets = {
        "MISSING_SIGNATURE": "No signature language was detected.",
        "NO_TERMINATION_CLAUSE": "No termination clause keywords were found.",
        "NO_GOVERNING_LAW": "No governing law or jurisdiction clause was detected.",
        "AUTO_RENEWAL": "Auto-renewal language was found.",
        "LOW_EXTRACTION_QUALITY": "The extracted text was too short or low quality.",
    }
    return {
        "rule_code": rule_code,
        "severity": severity,
        "snippet": snippet or snippets.get(rule_code),
        "suggestion": suggestions[rule_code],
    }


def _write_sample_file(*, user_id: str, filename: str, paragraphs: list[str]) -> tuple[Path, int, str]:
    storage_path = Path(settings.upload_dir) / user_id / f"seed-{filename}"
    _ensure_docx_file(storage_path, paragraphs)
    return storage_path, storage_path.stat().st_size, _sha256(storage_path)


def _ensure_docx_file(path: Path, paragraphs: list[str]) -> None:
    if path.exists():
        return
    path.parent.mkdir(parents=True, exist_ok=True)
    docx = DocxDocument()
    for paragraph in paragraphs:
        docx.add_paragraph(paragraph)
    docx.save(path)


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as input_file:
        for chunk in iter(lambda: input_file.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _add_seed_audit_logs(
    db,
    *,
    document: Document,
    actor_id,
    processed_at: datetime | None,
    review_status: str,
) -> None:
    db.add(
        AuditLog(
            action="document.uploaded",
            target_type="document",
            target_id=document.id,
            actor_id=actor_id,
            payload={"filename": document.filename, "seeded": True},
        )
    )
    if document.processing_status == "processing":
        db.add(
            AuditLog(
                action="document.processing.started",
                target_type="document",
                target_id=document.id,
                actor_id=actor_id,
                payload={"filename": document.filename, "seeded": True},
            )
        )
        return
    if document.processing_status == "failed":
        db.add(
            AuditLog(
                action="document.processing.failed",
                target_type="document",
                target_id=document.id,
                actor_id=actor_id,
                payload={"filename": document.filename, "seeded": True},
            )
        )
        return
    db.add(
        AuditLog(
            action="ai.review.completed",
            target_type="document",
            target_id=document.id,
            actor_id=actor_id,
            payload={
                "provider": "seed",
                "review_status": review_status,
                "processed_at": processed_at.isoformat() if processed_at else None,
                "seeded": True,
            },
        )
    )


if __name__ == "__main__":
    seed_users()
    if settings.seed_sample_documents:
        seed_sample_documents()
        print("Seed users and sample documents ready.")
    else:
        print("Seed users ready. Sample documents skipped.")
